#!/usr/bin/env python3
"""
Update Part II of De_xuat_giai_phap_ky_thuat_TCB_AIOps.docx
with content from Confluence DAB1 Document pages.

- Replaces PHẦN II section with Confluence content
- Applies formatting standards (Times New Roman, proper heading hierarchy)
- Adds review tracking (green highlight = new, yellow = modified)
- Generates REVIEW and clean versions
"""

import copy
import json
import os
import re
import sys
from datetime import datetime
from html.parser import HTMLParser
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Error: python-docx not found. Install: pip3 install python-docx")
    sys.exit(1)

# ─── Configuration ───────────────────────────────────────────────────────────
AUTHOR = "Khoa Than"
MODEL = "Opus 4.6"
PHASE = "Phase 1"
TIMESTAMP = datetime.now().strftime("%Y-%m-%d %H:%M")

INPUT_DOCX = "input/De_xuat_giai_phap_ky_thuat_TCB_AIOps.docx"
OUTPUT_REVIEW = "output/De_xuat_giai_phap_ky_thuat_TCB_AIOps_REVIEW.docx"
OUTPUT_CLEAN = "output/De_xuat_giai_phap_ky_thuat_TCB_AIOps.docx"
CONFLUENCE_DIR = "input/confluence"
MANIFEST_FILE = os.path.join(CONFLUENCE_DIR, "manifest.json")

# Highlight colors (WdColorIndex values for review tracking)
GREEN_HIGHLIGHT = 4    # wdBrightGreen - NEW content
YELLOW_HIGHLIGHT = 7   # wdYellow - MODIFIED content

# Formatting constants
FONT_BODY = "Times New Roman"
FONT_SIZE_BODY = Pt(13)
FONT_SIZE_H1 = Pt(16)
FONT_SIZE_H2 = Pt(14)
FONT_SIZE_H3 = Pt(13)
FONT_SIZE_TABLE = Pt(11)

# Confluence pages that map to design sections (order matters)
DESIGN_PAGES = [
    ("1412727495", "1. Introduction"),
    ("1411383773", "2. Requirements"),
    ("1413578790", "3. Current Solution"),
    ("1412727506", "4. Solution Design"),
    ("1412727525", "5. Infrastructure Design"),
    ("1413578820", "6. Security Design"),
    ("1412727545", "7. Deployment Design"),
]

APPENDIX_PAGES = [
    ("1411351001", "Appendix 1 - Business Outcome & Capability Mapping"),
    ("1412727628", "Appendix 2 - Non-Functional Requirements"),
    ("1413579036", "Appendix 3 - Synaptix Assistant"),
    ("1413611639", "Appendix 4 - Jira Incident Processing Sequence"),
    ("1413611652", "Appendix 5 - Design Decision Log"),
    ("1414004764", "Appendix 7 - Database Design"),
]


# ─── HTML to structured content parser ────────────────────────────────────────

class ConfluenceHTMLParser(HTMLParser):
    """Parse Confluence storage HTML into structured content blocks."""

    def __init__(self):
        super().__init__()
        self.blocks = []  # List of content blocks
        self.current_text = ""
        self.tag_stack = []
        self.in_table = False
        self.current_table = []
        self.current_row = []
        self.current_cell = ""
        self.current_cell_is_header = False
        self.list_stack = []  # Track nested lists
        self.skip_tags = {"ac:link", "ri:user", "ri:page", "ac:link-body",
                          "ac:image", "ri:attachment", "ac:structured-macro",
                          "ac:parameter", "ac:plain-text-body", "ac:rich-text-body",
                          "time", "colgroup", "col"}
        self.heading_level = 0
        self.in_bold = False
        self.in_italic = False
        self.in_code = False
        self.img_refs = []

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        self.tag_stack.append(tag)

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._flush_text()
            self.heading_level = int(tag[1])
        elif tag == "p" and not self.in_table:
            pass  # Will flush on endtag
        elif tag == "table":
            self._flush_text()
            self.in_table = True
            self.current_table = []
        elif tag == "tr":
            self.current_row = []
        elif tag in ("td", "th"):
            self.current_cell = ""
            self.current_cell_is_header = (tag == "th")
        elif tag in ("ul", "ol"):
            self._flush_text()
            self.list_stack.append(tag)
        elif tag == "li":
            pass
        elif tag == "strong" or tag == "b":
            self.in_bold = True
        elif tag == "em" or tag == "i":
            self.in_italic = True
        elif tag == "code":
            self.in_code = True
        elif tag == "br":
            if self.in_table:
                self.current_cell += "\n"
            else:
                self.current_text += "\n"
        elif tag == "img":
            src = attrs_dict.get("src", "")
            alt = attrs_dict.get("alt", "")
            if src:
                self.img_refs.append({"src": src, "alt": alt})
                placeholder = "{{IMAGE:" + (alt or "diagram") + "}}"
                if self.in_table:
                    self.current_cell += placeholder
                else:
                    self.current_text += placeholder
        elif tag == "a":
            pass  # Handle link text normally

    def handle_endtag(self, tag):
        if self.tag_stack and self.tag_stack[-1] == tag:
            self.tag_stack.pop()

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            text = self.current_text.strip()
            if text:
                self.blocks.append({
                    "type": "heading",
                    "level": self.heading_level,
                    "text": text
                })
            self.current_text = ""
            self.heading_level = 0
        elif tag == "p":
            if self.in_table:
                # Cell content
                pass
            elif self.list_stack:
                # List item text handled by li
                pass
            else:
                text = self.current_text.strip()
                if text:
                    self.blocks.append({
                        "type": "paragraph",
                        "text": text
                    })
                self.current_text = ""
        elif tag == "table":
            if self.current_table:
                self.blocks.append({
                    "type": "table",
                    "rows": self.current_table
                })
            self.in_table = False
            self.current_table = []
        elif tag == "tr":
            if self.current_row:
                self.current_table.append(self.current_row)
            self.current_row = []
        elif tag in ("td", "th"):
            self.current_row.append({
                "text": self.current_cell.strip(),
                "is_header": self.current_cell_is_header
            })
            self.current_cell = ""
        elif tag == "li":
            text = self.current_text.strip()
            if text:
                depth = len(self.list_stack) - 1
                list_type = self.list_stack[-1] if self.list_stack else "ul"
                self.blocks.append({
                    "type": "list_item",
                    "text": text,
                    "depth": depth,
                    "ordered": list_type == "ol"
                })
            self.current_text = ""
        elif tag in ("ul", "ol"):
            if self.list_stack:
                self.list_stack.pop()
        elif tag in ("strong", "b"):
            self.in_bold = False
        elif tag in ("em", "i"):
            self.in_italic = False
        elif tag == "code":
            self.in_code = False

    def handle_data(self, data):
        text = data
        if self.in_table:
            self.current_cell += text
        else:
            self.current_text += text

    def _flush_text(self):
        text = self.current_text.strip()
        if text:
            self.blocks.append({
                "type": "paragraph",
                "text": text
            })
        self.current_text = ""

    def get_blocks(self):
        self._flush_text()
        return self.blocks

    def get_images(self):
        return self.img_refs


def parse_confluence_html(html_content):
    """Parse Confluence HTML and return structured blocks."""
    parser = ConfluenceHTMLParser()
    # Remove XML comments
    html_content = re.sub(r'<!--.*?-->', '', html_content, flags=re.DOTALL)
    # Remove Confluence-specific XML namespaces
    html_content = re.sub(r'<ac:.*?/>', '', html_content, flags=re.DOTALL)
    html_content = re.sub(r'<ri:.*?/>', '', html_content, flags=re.DOTALL)
    try:
        parser.feed(html_content)
    except Exception as e:
        print(f"  Warning: HTML parse error: {e}")
    return parser.get_blocks(), parser.get_images()


# ─── Document manipulation ────────────────────────────────────────────────────

def find_section_boundaries(doc, start_text, end_text):
    """Find paragraph indices for section start and end."""
    start_idx = None
    end_idx = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if start_text in text and start_idx is None:
            start_idx = i
        elif end_text in text and start_idx is not None:
            end_idx = i
            break

    return start_idx, end_idx


def delete_paragraphs_between(doc, start_idx, end_idx):
    """Delete paragraphs between start (exclusive) and end (exclusive)."""
    # We need to work with the XML elements directly
    body = doc.element.body
    paragraphs = list(body.iterchildren(qn('w:p')))
    tables = list(body.iterchildren(qn('w:tbl')))

    # Get all body children in order
    children = list(body)

    # Find the XML elements for start and end paragraphs
    start_elem = doc.paragraphs[start_idx]._element
    end_elem = doc.paragraphs[end_idx]._element

    # Find positions in children
    start_pos = None
    end_pos = None
    for i, child in enumerate(children):
        if child is start_elem:
            start_pos = i
        elif child is end_elem:
            end_pos = i
            break

    if start_pos is None or end_pos is None:
        print(f"Error: Could not find section boundaries in XML")
        return

    # Remove all elements between start and end (exclusive)
    elements_to_remove = children[start_pos + 1:end_pos]
    for elem in elements_to_remove:
        body.remove(elem)

    print(f"  Removed {len(elements_to_remove)} elements between sections")


def apply_highlight(run, color_index):
    """Apply highlight color to a run."""
    rPr = run._element.get_or_add_rPr()
    highlight = parse_xml(f'<w:highlight {nsdecls("w")} w:val="{color_index}"/>')
    # Remove existing highlight
    for existing in rPr.findall(qn('w:highlight')):
        rPr.remove(existing)
    rPr.append(highlight)


def set_run_font(run, font_name=FONT_BODY, font_size=FONT_SIZE_BODY,
                 bold=False, italic=False, color=None):
    """Set font properties on a run."""
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # Ensure East Asian font is also set
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


def add_heading_paragraph(doc, text, level, after_element, is_new=True):
    """Add a heading paragraph after a specific element."""
    body = doc.element.body

    # Create paragraph with heading style
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:pStyle w:val="Heading{level}"/></w:pPr></w:p>')

    # Add run with text
    run_elem = parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{_escape_xml(text)}</w:t></w:r>')

    # Set font properties on the run
    rPr = parse_xml(f'<w:rPr {nsdecls("w")}>'
                     f'<w:rFonts w:ascii="{FONT_BODY}" w:hAnsi="{FONT_BODY}" w:eastAsia="{FONT_BODY}"/>'
                     f'<w:b/>'
                     f'</w:rPr>')

    if level == 1:
        sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="{int(FONT_SIZE_H1.pt * 2)}"/>')
        rPr.append(sz)
    elif level == 2:
        sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="{int(FONT_SIZE_H2.pt * 2)}"/>')
        rPr.append(sz)
    elif level >= 3:
        sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="{int(FONT_SIZE_H3.pt * 2)}"/>')
        rPr.append(sz)

    # Add highlight for review tracking
    if is_new:
        highlight = parse_xml(f'<w:highlight {nsdecls("w")} w:val="4"/>')  # Green
        rPr.append(highlight)

    run_elem.insert(0, rPr)
    p.append(run_elem)

    # Insert after the specified element
    after_element.addnext(p)
    return p


def add_paragraph(doc, text, after_element, is_new=True, bold=False, italic=False):
    """Add a paragraph after a specific element."""
    p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')

    # Set paragraph properties (spacing)
    pPr = parse_xml(f'<w:pPr {nsdecls("w")}>'
                     f'<w:spacing w:after="120" w:line="360" w:lineRule="auto"/>'
                     f'</w:pPr>')
    p.append(pPr)

    # Add run
    run_elem = parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{_escape_xml(text)}</w:t></w:r>')

    rPr_parts = [f'<w:rFonts w:ascii="{FONT_BODY}" w:hAnsi="{FONT_BODY}" w:eastAsia="{FONT_BODY}"/>',
                 f'<w:sz w:val="{int(FONT_SIZE_BODY.pt * 2)}"/>']
    if bold:
        rPr_parts.append('<w:b/>')
    if italic:
        rPr_parts.append('<w:i/>')
    if is_new:
        rPr_parts.append('<w:highlight w:val="4"/>')  # Green

    rPr = parse_xml(f'<w:rPr {nsdecls("w")}>{"".join(rPr_parts)}</w:rPr>')
    run_elem.insert(0, rPr)
    p.append(run_elem)

    after_element.addnext(p)
    return p


def add_list_item(doc, text, after_element, depth=0, ordered=False, is_new=True):
    """Add a list item paragraph."""
    p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')

    indent = 720 * (depth + 1)  # 720 twips = 0.5 inch per level
    pPr = parse_xml(f'<w:pPr {nsdecls("w")}>'
                     f'<w:pStyle w:val="ListParagraph"/>'
                     f'<w:ind w:left="{indent}" w:hanging="360"/>'
                     f'<w:spacing w:after="60" w:line="360" w:lineRule="auto"/>'
                     f'</w:pPr>')
    p.append(pPr)

    # Add bullet/number prefix
    prefix = "- " if not ordered else ""

    run_elem = parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{_escape_xml(prefix + text)}</w:t></w:r>')
    rPr_parts = [f'<w:rFonts w:ascii="{FONT_BODY}" w:hAnsi="{FONT_BODY}" w:eastAsia="{FONT_BODY}"/>',
                 f'<w:sz w:val="{int(FONT_SIZE_BODY.pt * 2)}"/>']
    if is_new:
        rPr_parts.append('<w:highlight w:val="4"/>')

    rPr = parse_xml(f'<w:rPr {nsdecls("w")}>{"".join(rPr_parts)}</w:rPr>')
    run_elem.insert(0, rPr)
    p.append(run_elem)

    after_element.addnext(p)
    return p


def add_table(doc, rows, after_element, is_new=True):
    """Add a table after a specific element."""
    if not rows:
        return after_element

    num_cols = max(len(row) for row in rows) if rows else 1

    # Create table XML
    tbl = parse_xml(f'<w:tbl {nsdecls("w")}></w:tbl>')

    # Table properties
    tblPr = parse_xml(
        f'<w:tblPr {nsdecls("w")}>'
        f'<w:tblStyle w:val="TableGrid"/>'
        f'<w:tblW w:w="0" w:type="auto"/>'
        f'<w:tblBorders>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
        f'<w:tblLook w:val="04A0" w:firstRow="1" w:lastRow="0" w:firstColumn="1" w:lastColumn="0" w:noHBand="0" w:noVBand="1"/>'
        f'</w:tblPr>'
    )
    tbl.append(tblPr)

    for row_idx, row in enumerate(rows):
        tr = parse_xml(f'<w:tr {nsdecls("w")}></w:tr>')

        for col_idx in range(num_cols):
            cell = row[col_idx] if col_idx < len(row) else {"text": "", "is_header": False}
            cell_text = cell.get("text", "")
            is_header = cell.get("is_header", False) or row_idx == 0

            tc = parse_xml(f'<w:tc {nsdecls("w")}></w:tc>')

            # Cell properties
            tcPr_parts = ['<w:tcW w:w="0" w:type="auto"/>']
            if is_header:
                tcPr_parts.append(
                    '<w:shd w:val="clear" w:color="auto" w:fill="1F4E79"/>'
                )
            tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}>{"".join(tcPr_parts)}</w:tcPr>')
            tc.append(tcPr)

            # Cell paragraph
            p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')

            # Truncate very long cell content
            if len(cell_text) > 500:
                cell_text = cell_text[:497] + "..."

            run = parse_xml(
                f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{_escape_xml(cell_text)}</w:t></w:r>'
            )

            rPr_parts = [
                f'<w:rFonts w:ascii="{FONT_BODY}" w:hAnsi="{FONT_BODY}" w:eastAsia="{FONT_BODY}"/>',
                f'<w:sz w:val="{int(FONT_SIZE_TABLE.pt * 2)}"/>'
            ]
            if is_header:
                rPr_parts.append('<w:b/>')
                rPr_parts.append('<w:color w:val="FFFFFF"/>')
            if is_new:
                rPr_parts.append('<w:highlight w:val="4"/>')

            rPr = parse_xml(f'<w:rPr {nsdecls("w")}>{"".join(rPr_parts)}</w:rPr>')
            run.insert(0, rPr)
            p.append(run)
            tc.append(p)
            tr.append(tc)

        tbl.append(tr)

    after_element.addnext(tbl)
    return tbl


def add_section_marker(doc, section_name, after_element):
    """Add a review tracking section marker (italic gray text)."""
    p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')

    pPr = parse_xml(f'<w:pPr {nsdecls("w")}>'
                     f'<w:spacing w:after="60"/>'
                     f'</w:pPr>')
    p.append(pPr)

    marker_text = f"Updated by {AUTHOR} using {MODEL} on {TIMESTAMP} | {PHASE} | Source: Confluence DAB1/{section_name}"
    run = parse_xml(
        f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{_escape_xml(marker_text)}</w:t></w:r>'
    )
    rPr = parse_xml(
        f'<w:rPr {nsdecls("w")}>'
        f'<w:rFonts w:ascii="{FONT_BODY}" w:hAnsi="{FONT_BODY}"/>'
        f'<w:sz w:val="18"/>'
        f'<w:i/>'
        f'<w:color w:val="808080"/>'
        f'</w:rPr>'
    )
    run.insert(0, rPr)
    p.append(run)

    after_element.addnext(p)
    return p


def _escape_xml(text):
    """Escape special XML characters."""
    if not text:
        return ""
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    text = text.replace('"', "&quot;")
    text = text.replace("'", "&apos;")
    return text


# ─── Main workflow ────────────────────────────────────────────────────────────

def load_manifest():
    """Load the Confluence manifest."""
    with open(MANIFEST_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def load_page_content(page_id):
    """Load and parse a Confluence page HTML file."""
    manifest = load_manifest()
    for page in manifest["pages"]:
        if page["id"] == page_id:
            html_path = os.path.join(CONFLUENCE_DIR, page["html_file"])
            if os.path.exists(html_path):
                with open(html_path, "r", encoding="utf-8") as f:
                    html = f.read()
                return parse_confluence_html(html)
    return [], []


def insert_confluence_content(doc, start_element, page_id, section_title):
    """Insert content from a Confluence page after the given element."""
    blocks, images = load_page_content(page_id)

    if not blocks:
        print(f"  Warning: No content found for page {page_id} ({section_title})")
        # Add placeholder paragraph
        return add_paragraph(doc, f"[Content pending for {section_title}]",
                           start_element, is_new=True, italic=True)

    print(f"  Inserting {len(blocks)} blocks for {section_title} "
          f"({len(images)} image refs)")

    current_elem = start_element
    block_count = 0

    for block in blocks:
        block_type = block["type"]

        if block_type == "heading":
            level = block["level"]
            text = block["text"]
            # Map Confluence heading levels to document levels
            # Confluence h1/h2/h3 -> doc level 2/3/4 (under Part II heading)
            doc_level = min(level + 1, 5)
            current_elem = add_heading_paragraph(doc, text, doc_level,
                                                 current_elem, is_new=True)
            block_count += 1

        elif block_type == "paragraph":
            text = block["text"]
            if text:
                # Check if it's a bold paragraph (sub-heading style)
                is_bold = text.startswith("**") and text.endswith("**")
                if is_bold:
                    text = text.strip("*").strip()
                current_elem = add_paragraph(doc, text, current_elem,
                                            is_new=True, bold=is_bold)
                block_count += 1

        elif block_type == "list_item":
            current_elem = add_list_item(doc, block["text"], current_elem,
                                        depth=block.get("depth", 0),
                                        ordered=block.get("ordered", False),
                                        is_new=True)
            block_count += 1

        elif block_type == "table":
            rows = block["rows"]
            if rows:
                current_elem = add_table(doc, rows, current_elem, is_new=True)
                block_count += 1

    print(f"  Inserted {block_count} elements")
    return current_elem


def add_change_log_page(doc, after_element, changes):
    """Add a change log page at the end of Part II."""
    current = add_heading_paragraph(doc, "Change Log", 2, after_element, is_new=True)
    current = add_section_marker(doc, "Change Log", current)

    summary_text = f"Total: {len(changes)} sections replaced with Confluence DAB1 content"
    current = add_paragraph(doc, summary_text, current, is_new=True, bold=True)

    # Build change log table
    rows = [
        [{"text": "Section", "is_header": True},
         {"text": "Change Type", "is_header": True},
         {"text": "Description", "is_header": True},
         {"text": "Author", "is_header": True},
         {"text": "Model", "is_header": True},
         {"text": "Date", "is_header": True}]
    ]

    for change in changes:
        rows.append([
            {"text": change["section"], "is_header": False},
            {"text": change["type"], "is_header": False},
            {"text": change["description"], "is_header": False},
            {"text": AUTHOR, "is_header": False},
            {"text": MODEL, "is_header": False},
            {"text": TIMESTAMP, "is_header": False},
        ])

    current = add_table(doc, rows, current, is_new=True)
    return current


def create_clean_version(review_path, clean_path):
    """Create clean version by removing highlights and tracking markers."""
    doc = Document(review_path)

    for para in doc.paragraphs:
        # Remove section markers (italic gray text with "Updated by")
        if para.text.strip().startswith("Updated by") and "Phase" in para.text:
            para.clear()
            continue

        # Remove highlights from runs
        for run in para.runs:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                for highlight in rPr.findall(qn('w:highlight')):
                    rPr.remove(highlight)

    # Remove highlights from table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        rPr = run._element.find(qn('w:rPr'))
                        if rPr is not None:
                            for highlight in rPr.findall(qn('w:highlight')):
                                rPr.remove(highlight)

    # Remove change log section if it exists
    # (This is a simplified approach - remove paragraphs with "Change Log" heading)

    doc.save(clean_path)
    print(f"  Clean version saved: {clean_path}")


def main():
    print("=" * 70)
    print("UPDATE PART II FROM CONFLUENCE DAB1 DOCUMENT")
    print("=" * 70)
    print(f"Author: {AUTHOR} | Model: {MODEL} | Phase: {PHASE}")
    print(f"Timestamp: {TIMESTAMP}")
    print()

    # Verify input files exist
    if not os.path.exists(INPUT_DOCX):
        print(f"Error: Input file not found: {INPUT_DOCX}")
        sys.exit(1)
    if not os.path.exists(MANIFEST_FILE):
        print(f"Error: Manifest not found: {MANIFEST_FILE}")
        sys.exit(1)

    # Load document
    print(f"Loading document: {INPUT_DOCX}")
    doc = Document(INPUT_DOCX)
    print(f"  Total paragraphs: {len(doc.paragraphs)}")

    # Find Part II boundaries
    print("\nFinding Part II section boundaries...")
    start_idx, end_idx = find_section_boundaries(
        doc, "DESIGN", "PHẦN III"
    )

    if start_idx is None:
        print("Error: Could not find Part II section start")
        sys.exit(1)
    if end_idx is None:
        print("Error: Could not find Part III section start")
        sys.exit(1)

    print(f"  Part II starts at paragraph {start_idx}: "
          f"{doc.paragraphs[start_idx].text[:80]}")
    print(f"  Part III starts at paragraph {end_idx}: "
          f"{doc.paragraphs[end_idx].text[:80]}")
    print(f"  Paragraphs to replace: {end_idx - start_idx - 1}")

    # Delete existing Part II content (keep the heading)
    print("\nRemoving existing Part II content...")
    delete_paragraphs_between(doc, start_idx, end_idx)

    # Get the Part II heading element to insert after
    part2_heading = doc.paragraphs[start_idx]._element

    # Add section marker
    current_elem = add_section_marker(doc, "Part II Overview", part2_heading)

    # Track changes for change log
    changes = []

    # Insert content from each Confluence design page
    print("\nInserting Confluence content...")

    for page_id, section_title in DESIGN_PAGES:
        print(f"\n--- Processing: {section_title} ---")

        # Add top-level section heading
        current_elem = add_heading_paragraph(
            doc, section_title, 2, current_elem, is_new=True
        )

        # Add section marker for review tracking
        current_elem = add_section_marker(doc, section_title, current_elem)

        # Insert page content
        current_elem = insert_confluence_content(
            doc, current_elem, page_id, section_title
        )

        changes.append({
            "section": section_title,
            "type": "NEW",
            "description": f"Replaced with content from Confluence DAB1 page {page_id}"
        })

    # Insert appendices
    print("\n\n--- Processing Appendices ---")
    current_elem = add_heading_paragraph(
        doc, "APPENDICES", 2, current_elem, is_new=True
    )

    for page_id, appendix_title in APPENDIX_PAGES:
        print(f"\n--- Processing: {appendix_title} ---")
        current_elem = add_heading_paragraph(
            doc, appendix_title, 3, current_elem, is_new=True
        )
        current_elem = add_section_marker(doc, appendix_title, current_elem)
        current_elem = insert_confluence_content(
            doc, current_elem, page_id, appendix_title
        )

        changes.append({
            "section": appendix_title,
            "type": "NEW",
            "description": f"Added from Confluence DAB1 page {page_id}"
        })

    # Add change log page
    print("\nAdding change log...")
    add_change_log_page(doc, current_elem, changes)

    # Save REVIEW version
    os.makedirs(os.path.dirname(OUTPUT_REVIEW), exist_ok=True)
    print(f"\nSaving REVIEW version: {OUTPUT_REVIEW}")
    doc.save(OUTPUT_REVIEW)
    print("  REVIEW version saved successfully!")

    # Create clean version
    print(f"\nCreating clean version: {OUTPUT_CLEAN}")
    create_clean_version(OUTPUT_REVIEW, OUTPUT_CLEAN)

    # Summary
    print("\n" + "=" * 70)
    print("SUMMARY")
    print("=" * 70)
    print(f"Sections replaced: {len(changes)}")
    print(f"REVIEW version: {OUTPUT_REVIEW}")
    print(f"Clean version: {OUTPUT_CLEAN}")
    print(f"\nReview tracking:")
    print(f"  GREEN highlight = NEW content from Confluence")
    print(f"  Section markers = Author/Model/Date/Phase tracking")
    print(f"  Change Log = Last page of Part II")
    print()
    for change in changes:
        print(f"  [{change['type']}] {change['section']}")


if __name__ == "__main__":
    main()
