#!/usr/bin/env python3
"""
Format Word (.docx) files to match formatting standards and apply review tracking.

This script applies comprehensive formatting to Word documents including:
- Typography (fonts, sizes, styles)
- Layout (margins, spacing, page setup)
- Tables (borders, colors, formatting)
- Headers and footers (page numbers, document title)
- Review tracking (comments, highlights, change log)

Usage:
    python scripts/word_formatter.py output/proposal.docx
    python scripts/word_formatter.py output/proposal.docx --no-review-mode
    python scripts/word_formatter.py output/proposal.docx --config config.yaml --changes changes.json
"""

import argparse
import json
import logging
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

try:
    import yaml
except ImportError:
    yaml = None
    logging.warning("PyYAML not available. config.yaml will not be read.")

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    Document = None
    logging.error("python-docx is required. Install with: pip install python-docx")


logging.basicConfig(
    level=logging.INFO,
    format='%(levelname)s: %(message)s'
)
logger = logging.getLogger(__name__)


# Color constants
COLOR_NEW_HIGHLIGHT = RGBColor(226, 239, 218)  # #E2EFDA - Light green
COLOR_MODIFIED_HIGHLIGHT = RGBColor(255, 242, 204)  # #FFF2CC - Yellow
COLOR_TABLE_HEADER_BG = RGBColor(31, 78, 121)  # #1F4E79 - Dark blue
COLOR_TABLE_HEADER_TEXT = RGBColor(255, 255, 255)  # White
COLOR_TABLE_ALT_ROW = RGBColor(242, 242, 242)  # #F2F2F2 - Light gray
COLOR_UPDATE_MARKER = RGBColor(128, 128, 128)  # Gray for update markers


def load_config(config_path: Optional[Path] = None) -> Dict:
    """
    Load configuration from config.yaml file.
    
    Args:
        config_path: Optional path to config file. If None, looks for config.yaml in project root.
        
    Returns:
        Dictionary with config values, or empty dict if not found.
    """
    if yaml is None:
        logger.warning("PyYAML not available. Skipping config.yaml.")
        return {}
    
    if config_path is None:
        # Look for config.yaml in project root (parent of scripts/)
        project_root = Path(__file__).parent.parent
        config_path = project_root / "config.yaml"
    
    config_path = Path(config_path)
    
    if not config_path.exists():
        logger.warning(f"Config file not found: {config_path}. Using defaults.")
        return {}
    
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            config = yaml.safe_load(f) or {}
        logger.info(f"Loaded config from {config_path}")
        return config
    except Exception as e:
        logger.warning(f"Error loading config.yaml: {e}. Using defaults.")
        return {}


def load_changes(changes_path: Optional[Path] = None) -> List[Dict]:
    """
    Load changes manifest from JSON file.
    
    Args:
        changes_path: Path to changes.json file.
        
    Returns:
        List of change dictionaries, or empty list if not found.
    """
    if changes_path is None:
        logger.warning("No changes.json provided. Review tracking will be limited.")
        return []
    
    changes_path = Path(changes_path)
    
    if not changes_path.exists():
        logger.warning(f"Changes file not found: {changes_path}")
        return []
    
    try:
        with open(changes_path, 'r', encoding='utf-8') as f:
            changes = json.load(f)
        logger.info(f"Loaded {len(changes)} change entries from {changes_path}")
        return changes
    except Exception as e:
        logger.error(f"Error loading changes.json: {e}")
        return []


def set_document_properties(doc: Document, config: Dict):
    """
    Set document core properties from config.
    
    Args:
        doc: Document object
        config: Configuration dictionary
    """
    author_name = config.get('author', {}).get('name', 'Unknown')
    
    # Set author
    doc.core_properties.author = author_name
    
    # Set last modified by (include model if available)
    model = config.get('models', {}).get('phase1_create', 'Unknown Model')
    doc.core_properties.last_modified_by = f"{author_name} / {model}"
    
    # Set modified timestamp
    doc.core_properties.modified = datetime.now()
    
    # Increment revision if it exists
    try:
        if doc.core_properties.revision:
            doc.core_properties.revision += 1
        else:
            doc.core_properties.revision = 1
    except:
        pass


def set_margins(doc: Document):
    """
    Set document margins: Left 3cm, Right 2cm, Top 2cm, Bottom 2cm.
    
    Args:
        doc: Document object
    """
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.page_height = Cm(29.7)  # A4 height
        section.page_width = Cm(21.0)   # A4 width


def format_paragraph_fonts(doc: Document):
    """
    Format paragraph fonts according to standards.
    
    Args:
        doc: Document object
    """
    for paragraph in doc.paragraphs:
        style = paragraph.style
        
        # Body text: Times New Roman, 13pt
        if style.name == 'Normal' or style.name.startswith('Body'):
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
        
        # Heading 1: Times New Roman Bold, 16pt
        elif style.name == 'Heading 1':
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)
                run.font.bold = True
            # Set paragraph spacing
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(6)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # Heading 2: Times New Roman Bold, 14pt
        elif style.name == 'Heading 2':
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.bold = True
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(6)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # Heading 3: Times New Roman Bold Italic, 13pt
        elif style.name == 'Heading 3':
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.italic = True
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(6)
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # Default paragraph formatting
        else:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            paragraph_format.space_after = Pt(6)


def format_tables(doc: Document):
    """
    Format all tables with borders, header styling, and alternating row colors.
    
    Args:
        doc: Document object
    """
    for table in doc.tables:
        # Set table alignment
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Format header row (first row)
        if len(table.rows) > 0:
            header_row = table.rows[0]
            for cell in header_row.cells:
                # Set background color
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '1F4E79')
                cell._element.get_or_add_tcPr().append(shading_elm)
                
                # Format text in header cells
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        run.font.bold = True
                        run.font.color.rgb = COLOR_TABLE_HEADER_TEXT
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_after = Pt(0)
        
        # Format data rows
        for row_idx, row in enumerate(table.rows[1:], start=1):
            # Alternating row colors
            if row_idx % 2 == 0:
                for cell in row.cells:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'F2F2F2')
                    cell._element.get_or_add_tcPr().append(shading_elm)
            
            # Format cell text
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                    paragraph.paragraph_format.space_after = Pt(0)
        
        # Add borders to all cells
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                
                # Create borders
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcBorders.append(border)
                
                tcPr.append(tcBorders)


def add_page_numbers(doc: Document):
    """
    Add page numbers to footer (center-aligned).
    
    Args:
        doc: Document object
    """
    for section in doc.sections:
        footer = section.footer
        
        # Clear existing footer content
        for paragraph in footer.paragraphs:
            paragraph.clear()
        
        # Add page number paragraph
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number field
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        
        # Format footer font
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)


def add_document_title_header(doc: Document):
    """
    Add document title to header (right-aligned).
    
    Args:
        doc: Document object
    """
    # Get document title from core properties
    title = doc.core_properties.title or "Document"
    
    for section in doc.sections:
        header = section.header
        
        # Clear existing header content
        for paragraph in header.paragraphs:
            paragraph.clear()
        
        # Add title paragraph
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run = paragraph.add_run(title)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)


def add_comment_to_paragraph(doc: Document, paragraph, comment_text: str, author: str):
    """
    Add a Word comment to a paragraph using XML manipulation.
    
    Note: This creates comment markers in the document. Full comment functionality
    requires creating entries in comments.xml, which is complex. The markers will
    be visible but may need manual comment entry in Word for full functionality.
    
    Args:
        doc: Document object
        paragraph: Paragraph object
        comment_text: Text of the comment
        author: Author name for the comment
    """
    if Document is None:
        return
    
    try:
        # Generate a simple comment ID (incrementing from 0)
        # In a full implementation, we'd track this properly
        comment_id = '0'  # Simple ID - in production, track IDs properly
        
        # Create comment range start with ID
        comment_range_start = OxmlElement('w:commentRangeStart')
        comment_range_start.set(qn('w:id'), comment_id)
        
        # Create comment reference run (the visible comment marker)
        comment_ref_run = OxmlElement('w:r')
        comment_ref = OxmlElement('w:commentReference')
        comment_ref.set(qn('w:id'), comment_id)
        comment_ref_run.append(comment_ref)
        
        # Create comment range end
        comment_range_end = OxmlElement('w:commentRangeEnd')
        comment_range_end.set(qn('w:id'), comment_id)
        
        # Insert comment markers into paragraph
        # Wrap the paragraph content with comment markers
        p = paragraph._element
        
        # Insert comment range start at beginning
        if len(p) > 0:
            p.insert(0, comment_range_start)
        else:
            p.append(comment_range_start)
        
        # Insert comment reference (visible marker)
        p.append(comment_ref_run)
        
        # Insert comment range end at end
        p.append(comment_range_end)
        
        # Note: To create the actual comment content in comments.xml, we would need to:
        # 1. Get or create comments part
        # 2. Create a comment element with the text
        # 3. Link it to the comment ID
        # This is complex and requires deeper XML manipulation.
        # For now, the markers are added and can be manually filled in Word.
        
    except Exception as e:
        logger.warning(f"Could not add comment: {e}. Comment markers may not be fully functional.")


def apply_highlight(run, highlight_color: RGBColor):
    """
    Apply highlight color to a run.
    
    Args:
        run: Run object
        highlight_color: RGBColor for highlight
    """
    rPr = run._element.get_or_add_rPr()
    
    # Set highlight color
    highlight = OxmlElement('w:highlight')
    if highlight_color == COLOR_NEW_HIGHLIGHT:
        highlight.set(qn('w:val'), 'green')
    elif highlight_color == COLOR_MODIFIED_HIGHLIGHT:
        highlight.set(qn('w:val'), 'yellow')
    else:
        highlight.set(qn('w:val'), 'yellow')  # Default
    
    rPr.append(highlight)


def add_section_marker(doc: Document, heading_paragraph, author: str, model: str, phase: int):
    """
    Add section-level update marker after a heading paragraph.
    
    Args:
        doc: Document object
        heading_paragraph: The heading paragraph to add marker after
        author: Author name
        model: Model name
        phase: Phase number
    """
    # Get the parent element and insert after
    parent = heading_paragraph._element.getparent()
    marker_p_elem = OxmlElement('w:p')
    
    # Create paragraph properties for formatting
    pPr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '240')  # 6pt = 240 twips
    spacing.set(qn('w:before'), '0')
    pPr.append(spacing)
    marker_p_elem.append(pPr)
    
    # Create run with marker text
    date_str = datetime.now().strftime('%Y-%m-%d')
    marker_text = f"Updated by {author} using {model} on {date_str} | Phase {phase}"
    
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Font name
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    
    # Font size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')  # 9pt = 18 half-points
    rPr.append(sz)
    
    # Italic
    i = OxmlElement('w:i')
    rPr.append(i)
    
    # Color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '808080')  # Gray
    rPr.append(color)
    
    r.append(rPr)
    
    # Add text
    t = OxmlElement('w:t')
    t.text = marker_text
    r.append(t)
    
    marker_p_elem.append(r)
    
    # Insert after heading paragraph
    heading_idx = parent.index(heading_paragraph._element)
    parent.insert(heading_idx + 1, marker_p_elem)


def find_section_paragraph(doc: Document, section_name: str) -> Optional:
    """
    Find a paragraph that matches a section name (heading).
    
    Args:
        doc: Document object
        section_name: Section name to find
        
    Returns:
        Paragraph object if found, None otherwise
    """
    # Try exact match first
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == section_name:
            if paragraph.style.name.startswith('Heading'):
                return paragraph
    
    # Try partial match
    for paragraph in doc.paragraphs:
        if section_name in paragraph.text:
            if paragraph.style.name.startswith('Heading'):
                return paragraph
    
    return None


def create_change_log_page(doc: Document, changes: List[Dict], config: Dict):
    """
    Create a Change Log page at the end of the document.
    
    Args:
        doc: Document object
        changes: List of change dictionaries
        config: Configuration dictionary
    """
    # Add page break
    doc.add_page_break()
    
    # Add heading
    heading = doc.add_heading('Change Log', level=1)
    
    # Count changes by type
    change_counts = {}
    for change in changes:
        change_type = change.get('change_type', 'UNKNOWN')
        change_counts[change_type] = change_counts.get(change_type, 0) + 1
    
    # Add summary paragraph
    summary_parts = []
    for change_type, count in change_counts.items():
        summary_parts.append(f"{count} {change_type.lower()}")
    summary_text = f"Total: {', '.join(summary_parts)}"
    
    summary_para = doc.add_paragraph(summary_text)
    summary_para.style = 'Normal'
    summary_para.paragraph_format.space_after = Pt(12)
    
    # Create table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Section'
    header_cells[1].text = 'Change Type'
    header_cells[2].text = 'Description'
    header_cells[3].text = 'Author'
    header_cells[4].text = 'Model'
    header_cells[5].text = 'Date'
    
    # Add data rows
    for change in changes:
        row = table.add_row()
        row.cells[0].text = change.get('section', '')
        row.cells[1].text = change.get('change_type', '')
        row.cells[2].text = change.get('description', '')
        row.cells[3].text = change.get('author', config.get('author', {}).get('name', 'Unknown'))
        row.cells[4].text = change.get('model', 'Unknown')
        row.cells[5].text = change.get('date', datetime.now().strftime('%Y-%m-%d'))
    
    # Format the table
    format_tables(doc)


def apply_review_tracking(doc: Document, changes: List[Dict], config: Dict, review_mode: bool):
    """
    Apply review tracking to the document.
    
    Args:
        doc: Document object
        changes: List of change dictionaries
        config: Configuration dictionary
        review_mode: Whether review mode is enabled
    """
    if not review_mode or not changes:
        return
    
    author = config.get('author', {}).get('name', 'Unknown')
    model = config.get('models', {}).get('phase1_create', 'Unknown Model')
    
    # Group changes by section
    section_changes = {}
    for change in changes:
        section = change.get('section', '')
        if section not in section_changes:
            section_changes[section] = []
        section_changes[section].append(change)
    
    # Apply tracking for each section
    for section_name, section_change_list in section_changes.items():
        # Find the section paragraph
        section_para = find_section_paragraph(doc, section_name)
        if section_para is None:
            logger.warning(f"Section '{section_name}' not found in document")
            continue
        
        # Add section marker
        phase = section_change_list[0].get('phase', 1)
        add_section_marker(section_para, author, model, phase)
        
        # Find paragraphs in this section and apply highlights/comments
        # This is simplified - in practice, you'd need to identify all paragraphs
        # belonging to this section (until next heading)
        current_para = section_para
        section_started = False
        
        for change in section_change_list:
            change_type = change.get('change_type', '')
            description = change.get('description', '')
            
            # Find the next paragraph after the heading
            # In a real implementation, you'd track which paragraphs belong to which section
            # For now, we'll apply to the first paragraph after the heading
            para_idx = None
            for i, para in enumerate(doc.paragraphs):
                if para == current_para:
                    para_idx = i
                    break
            
            if para_idx is not None and para_idx + 1 < len(doc.paragraphs):
                target_para = doc.paragraphs[para_idx + 1]
                
                # Apply highlight based on change type
                if change_type == 'NEW':
                    for run in target_para.runs:
                        apply_highlight(run, COLOR_NEW_HIGHLIGHT)
                elif change_type == 'MODIFIED':
                    for run in target_para.runs:
                        apply_highlight(run, COLOR_MODIFIED_HIGHLIGHT)
                
                # Add comment
                comment_text = f"[PHASE {phase}] {author}/{model} - {description}"
                add_comment_to_paragraph(doc, target_para, comment_text, author)
    
    # Create change log page
    create_change_log_page(doc, changes, config)


def format_document(
    doc_path: Path,
    config: Optional[Dict] = None,
    changes: Optional[List[Dict]] = None,
    review_mode: bool = True
) -> None:
    """
    Format a Word document according to standards and apply review tracking.
    
    Args:
        doc_path: Path to the .docx file
        config: Configuration dictionary (optional)
        changes: List of change dictionaries (optional)
        review_mode: Whether to apply review tracking
    """
    if Document is None:
        raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    if config is None:
        config = {}
    if changes is None:
        changes = []
    
    logger.info(f"Formatting document: {doc_path}")
    
    # Load document
    doc = Document(str(doc_path))
    
    # Set document properties
    set_document_properties(doc, config)
    
    # Apply formatting
    set_margins(doc)
    format_paragraph_fonts(doc)
    format_tables(doc)
    add_page_numbers(doc)
    add_document_title_header(doc)
    
    # Apply review tracking if enabled
    if review_mode:
        apply_review_tracking(doc, changes, config, review_mode)
    
    # Save document
    doc.save(str(doc_path))
    logger.info(f"Document formatted and saved: {doc_path}")


def main():
    """Main entry point for the script."""
    parser = argparse.ArgumentParser(
        description='Format Word documents to match standards and apply review tracking',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python scripts/word_formatter.py output/proposal.docx
  python scripts/word_formatter.py output/proposal.docx --no-review-mode
  python scripts/word_formatter.py output/proposal.docx --config config.yaml --changes changes.json
        """
    )
    
    parser.add_argument(
        'doc_path',
        type=str,
        help='Path to the .docx file to format'
    )
    
    parser.add_argument(
        '--review-mode',
        action='store_true',
        default=True,
        help='Enable review tracking (default: enabled)'
    )
    
    parser.add_argument(
        '--no-review-mode',
        dest='review_mode',
        action='store_false',
        help='Disable review tracking'
    )
    
    parser.add_argument(
        '--config',
        type=str,
        default=None,
        help='Path to config.yaml file (default: looks for config.yaml in project root)'
    )
    
    parser.add_argument(
        '--changes',
        type=str,
        default=None,
        help='Path to changes.json file with change manifest'
    )
    
    parser.add_argument(
        '--verbose',
        action='store_true',
        help='Enable verbose logging'
    )
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # Validate input file
    doc_path = Path(args.doc_path)
    if not doc_path.exists():
        logger.error(f"Document not found: {doc_path}")
        sys.exit(1)
    
    if doc_path.suffix.lower() != '.docx':
        logger.error(f"File must be a .docx file: {doc_path}")
        sys.exit(1)
    
    # Load config
    config_path = Path(args.config) if args.config else None
    config = load_config(config_path)
    
    # Load changes
    changes_path = Path(args.changes) if args.changes else None
    changes = load_changes(changes_path)
    
    # Format document
    try:
        format_document(
            doc_path=doc_path,
            config=config,
            changes=changes,
            review_mode=args.review_mode
        )
        logger.info("Formatting completed successfully")
    except Exception as e:
        logger.error(f"Error formatting document: {e}", exc_info=args.verbose)
        sys.exit(1)


if __name__ == "__main__":
    main()
